"""Convert a markdown note to a .docx file (bytes). Reasonable markdown subset."""
import io
import re

from docx import Document
from docx.shared import Pt, RGBColor

_INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[\[[^\]]+?\]\])")

BODY_FONT = "Georgia"          # clean, readable, ships on every machine
HEAD_FONT = "Georgia"
CODE_FONT = "Consolas"


def _style_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("222222")
    pf = normal.paragraph_format
    pf.line_spacing = 1.3
    pf.space_after = Pt(8)

    heads = [(1, 21, "111111"), (2, 16, "1a1a1a"), (3, 13, "333333"), (4, 12, "333333")]
    for lvl, sz, col in heads:
        try:
            st = doc.styles[f"Heading {lvl}"]
        except KeyError:
            continue
        st.font.name = HEAD_FONT
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(col)
        st.paragraph_format.space_before = Pt(14)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.15
    try:
        t = doc.styles["Title"]
        t.font.name = HEAD_FONT
        t.font.size = Pt(28)
        t.font.bold = True
        t.font.color.rgb = RGBColor.from_string("111111")
    except KeyError:
        pass


def _add_inline(p, text):
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            p.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = CODE_FONT
            r.font.color.rgb = RGBColor.from_string("a03060")
        elif part.startswith("[[") and part.endswith("]]"):
            inner = part[2:-2].split("|")[-1].split("#")[0].strip()
            r = p.add_run(inner); r.italic = True   # wikilink → italic text
        else:
            p.add_run(part)


def md_to_docx(md: str, title: str = "") -> bytes:
    doc = Document()
    _style_doc(doc)
    if title:
        doc.add_heading(title, 0)

    in_code = False
    code = []
    for line in (md or "").split("\n"):
        if line.lstrip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code))
                r.font.name = CODE_FONT; r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor.from_string("3a3a3a")
                code = []; in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code.append(line)
            continue

        s = line.rstrip()
        if not s.strip():
            continue
        h = re.match(r"^(#{1,6})\s+(.*)", s)
        if h:
            doc.add_heading(h.group(2), min(len(h.group(1)), 4))
            continue
        if re.match(r"^[-*]\s+", s):
            _add_inline(doc.add_paragraph(style="List Bullet"), re.sub(r"^[-*]\s+", "", s))
            continue
        if re.match(r"^\d+\.\s+", s):
            _add_inline(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s+", "", s))
            continue
        if re.match(r"^[-_*]{3,}$", s):
            doc.add_paragraph("─" * 30)
            continue
        if s.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            _add_inline(p, s.lstrip("> ").rstrip())
            continue
        _add_inline(doc.add_paragraph(), s)

    if in_code and code:
        r = doc.add_paragraph().add_run("\n".join(code)); r.font.name = CODE_FONT

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
