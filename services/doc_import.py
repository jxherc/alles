"""
Import a document into the vault as markdown. Handles .md/.txt (passthrough),
.docx (python-docx), .html (stripped to md), .pdf (pypdf, graceful if missing).
"""
import io
import re
from html.parser import HTMLParser


def import_document(filename: str, data: bytes) -> dict:
    name = (filename or "imported").replace("\\", "/").rsplit("/", 1)[-1]
    stem, _, ext = name.rpartition(".")
    ext = ext.lower()
    stem = (stem or name).strip() or "imported"
    if ext in ("md", "markdown", "txt", ""):
        content = data.decode("utf-8", "replace")
    elif ext == "docx":
        content = _docx_to_md(data)
    elif ext in ("html", "htm"):
        content = _html_to_md(data.decode("utf-8", "replace"))
    elif ext == "pdf":
        content = _pdf_to_md(data)
    else:
        raise ValueError(f"can't import .{ext} — try .md, .txt, .docx, .html or .pdf")
    return {"name": stem, "content": (content or "").strip() + "\n"}


def _docx_to_md(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    out = []
    for p in doc.paragraphs:
        text = p.text.rstrip()
        style = ((p.style.name if p.style else "") or "").lower()
        if not text:
            out.append("")
            continue
        m = re.match(r"heading (\d)", style)
        if m:
            out.append("#" * min(int(m.group(1)), 6) + " " + text)
        elif style == "title":
            out.append("# " + text)
        elif "list number" in style:
            out.append("1. " + text)
        elif "list" in style:
            out.append("- " + text)
        elif style in ("quote", "intense quote"):
            out.append("> " + text)
        else:
            out.append(text)
    # tables come after the body (python-docx doesn't interleave them) — good enough
    for t in doc.tables:
        rows = ["| " + " | ".join(c.text.strip().replace("\n", " ") for c in r.cells) + " |" for r in t.rows]
        if rows:
            cols = rows[0].count("|") - 1
            rows.insert(1, "| " + " | ".join(["---"] * cols) + " |")
            out.append("")
            out.extend(rows)
    return "\n".join(out)


class _HtmlToMd(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.out = []
        self.skip = 0
        self.lists = []        # stack of [kind, counter]
        self.in_link = False
        self.href = None
        self.link_txt = []
        self.pre = 0

    def handle_starttag(self, tag, attrs):
        a = dict(attrs)
        if tag in ("script", "style", "head"):
            self.skip += 1
        elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self.out.append("\n\n" + "#" * int(tag[1]) + " ")
        elif tag == "p":
            self.out.append("\n\n")
        elif tag == "br":
            self.out.append("  \n")
        elif tag in ("strong", "b"):
            self.out.append("**")
        elif tag in ("em", "i"):
            self.out.append("*")
        elif tag == "code" and not self.pre:
            self.out.append("`")
        elif tag == "blockquote":
            self.out.append("\n\n> ")
        elif tag == "hr":
            self.out.append("\n\n---\n\n")
        elif tag == "ul":
            self.lists.append(["ul", 0])
        elif tag == "ol":
            self.lists.append(["ol", 0])
        elif tag == "li":
            indent = "  " * max(0, len(self.lists) - 1)
            if self.lists and self.lists[-1][0] == "ol":
                self.lists[-1][1] += 1
                self.out.append(f"\n{indent}{self.lists[-1][1]}. ")
            else:
                self.out.append(f"\n{indent}- ")
        elif tag == "a":
            self.in_link = True
            self.href = a.get("href")
            self.link_txt = []
        elif tag == "pre":
            self.pre += 1
            self.out.append("\n\n```\n")

    def handle_endtag(self, tag):
        if tag in ("script", "style", "head"):
            self.skip = max(0, self.skip - 1)
        elif tag in ("strong", "b"):
            self.out.append("**")
        elif tag in ("em", "i"):
            self.out.append("*")
        elif tag == "code" and not self.pre:
            self.out.append("`")
        elif tag in ("h1", "h2", "h3", "h4", "h5", "h6", "p", "blockquote"):
            self.out.append("\n")
        elif tag in ("ul", "ol"):
            if self.lists:
                self.lists.pop()
        elif tag == "a":
            txt = "".join(self.link_txt).strip()
            if self.href and txt:
                self.out.append(f"[{txt}]({self.href})")
            elif txt:
                self.out.append(txt)
            self.in_link = False
            self.href = None
            self.link_txt = []
        elif tag == "pre":
            self.pre = max(0, self.pre - 1)
            self.out.append("\n```\n")

    def handle_data(self, data):
        if self.skip:
            return
        if self.in_link:
            self.link_txt.append(data)
            return
        if self.pre:
            self.out.append(data)
            return
        self.out.append(re.sub(r"\s+", " ", data))

    def md(self):
        text = "".join(self.out)
        text = re.sub(r"[ \t]+\n", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


def _html_to_md(s: str) -> str:
    p = _HtmlToMd()
    p.feed(s)
    return p.md()


def _pdf_to_md(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception:
        try:
            from PyPDF2 import PdfReader   # older name, same api
        except Exception:
            raise ValueError("PDF import needs pypdf — run: pip install pypdf")
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for pg in reader.pages:
        try:
            txt = (pg.extract_text() or "").strip()
        except Exception:
            txt = ""
        if txt:
            pages.append(txt)
    return "\n\n".join(pages)
