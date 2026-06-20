"""6b UI verification — grid view, duplicates + activity smart folders, office preview. :8852.
Run with ALLES_DATA set so the seed files land in the server's files dir."""

import os
import shutil
import sys
from pathlib import Path

from playwright.sync_api import sync_playwright

FILES = "http://files.localhost:8852"
EVID = Path(__file__).resolve().parent.parent / "docs" / "evidence" / "6b"
IGNORE = ("Failed to load resource", "net::", "ERR_", "favicon", "401", "Load failed")


def _reseed():
    base = Path(os.environ["ALLES_DATA"]) / "files"
    seed = base / "browse6b"
    if seed.exists():
        shutil.rmtree(seed)
    seed.mkdir(parents=True, exist_ok=True)

    from PIL import Image

    Image.new("RGB", (16, 16), (80, 120, 200)).save(str(seed / "pic.png"))

    from docx import Document

    doc = Document()
    doc.add_paragraph("The quick brown fox")
    doc.add_paragraph("jumps over the lazy dog")
    doc.save(str(seed / "memo.docx"))

    (seed / "copy1.txt").write_text("duplicate me please")
    (seed / "copy2.txt").write_text("duplicate me please")
    (seed / "readme.txt").write_text("hello readme body")


def main():
    EVID.mkdir(parents=True, exist_ok=True)
    _reseed()
    r = {}
    errs = []
    with sync_playwright() as p:
        b = p.chromium.launch()
        pg = b.new_context().new_page()
        pg.on(
            "console",
            lambda m: (
                errs.append(m.text)
                if m.type == "error" and not any(x in m.text for x in IGNORE)
                else None
            ),
        )
        pg.on(
            "pageerror",
            lambda e: errs.append(str(e)) if not any(x in str(e) for x in IGNORE) else None,
        )

        # ---- land in the seeded folder ----
        pg.goto(f"{FILES}/?p=browse6b", wait_until="domcontentloaded")
        pg.wait_for_selector('.file-row[data-path="browse6b/pic.png"]', timeout=15000)

        # ---- grid toggle on ----
        pg.click("#files-view-toggle")
        pg.wait_for_selector("#files-list.files-grid", timeout=5000)
        r["grid_toggle"] = pg.query_selector("#files-list.files-grid") is not None

        # ---- grid shows a real image thumbnail ----
        r["grid_shows_thumb"] = (
            pg.query_selector('.file-row[data-path="browse6b/pic.png"] .file-thumb img') is not None
        )
        pg.screenshot(path=str(EVID / "files-grid.png"))

        # ---- toggle back to list ----
        pg.click("#files-view-toggle")
        pg.wait_for_timeout(300)
        r["list_toggle_back"] = pg.query_selector("#files-list.files-grid") is None

        # ---- docx preview shows extracted text ----
        pg.eval_on_selector(
            '.file-row[data-path="browse6b/memo.docx"] .file-name', "el => el.click()"
        )
        pg.wait_for_selector(".files-preview-doc", timeout=6000)
        r["docx_preview_text"] = "quick brown fox" in (pg.text_content(".files-preview-doc") or "")
        pg.screenshot(path=str(EVID / "files-docx-preview.png"))
        pg.click("#files-preview-close")

        # ---- plain preview pane opens ----
        pg.eval_on_selector(
            '.file-row[data-path="browse6b/readme.txt"] .file-name', "el => el.click()"
        )
        pg.wait_for_selector(".files-preview-pre", timeout=6000)
        r["preview_pane_opens"] = "hello readme body" in (
            pg.text_content(".files-preview-pre") or ""
        )
        pg.click("#files-preview-close")

        # ---- duplicates smart folder (root: smart bar is injected there) ----
        pg.goto(f"{FILES}/", wait_until="domcontentloaded")
        pg.wait_for_selector('.files-smart[data-kind="__duplicates"]', timeout=10000)
        pg.click('.files-smart[data-kind="__duplicates"]')
        pg.wait_for_selector(".files-dupgroup", timeout=6000)
        dup_txt = pg.text_content("#files-list") or ""
        r["duplicates_lists_group"] = (
            "browse6b/copy1.txt" in dup_txt and "browse6b/copy2.txt" in dup_txt
        )
        pg.screenshot(path=str(EVID / "files-duplicates.png"))

        # ---- activity smart folder ----
        pg.click("#files-smart-back")
        pg.wait_for_selector('.files-smart[data-kind="__activity"]', timeout=8000)
        pg.click('.files-smart[data-kind="__activity"]')
        # wait on an activity-specific row (root only has the "browse6b" dir row → avoid that race)
        pg.wait_for_selector('.file-row[data-path="browse6b/readme.txt"]', timeout=6000)
        act_txt = pg.text_content("#files-list") or ""
        r["activity_lists"] = "browse6b/readme.txt" in act_txt
        pg.screenshot(path=str(EVID / "files-activity.png"))

        r["zero_console_errors"] = len(errs) == 0
        b.close()

    ok = all(r.values())
    lines = [f"{'PASS' if v else 'FAIL'}  {k}" for k, v in r.items()]
    if errs:
        lines.append(f"console_errors: {errs[:8]}")
    out = "\n".join(lines)
    (EVID / "pw_files_browse_6b.txt").write_text(out, encoding="utf-8")
    print(out)
    print(f"\n{sum(bool(v) for v in r.values())}/{len(r)} assertions passed")
    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(main())
