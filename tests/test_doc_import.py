import io
import unittest

from services import doc_import


class DocImportTests(unittest.TestCase):
    def test_txt_passthrough(self):
        r = doc_import.import_document("my notes.txt", b"hello world")
        self.assertEqual(r["name"], "my notes")
        self.assertIn("hello world", r["content"])

    def test_md_passthrough(self):
        r = doc_import.import_document("a.md", b"# Title\n\n- x")
        self.assertIn("# Title", r["content"])

    def test_html_to_md(self):
        html = (
            "<html><head><style>.x{color:red}</style></head><body>"
            "<h1>Hi</h1><p>a <strong>bold</strong> word and "
            "<a href='http://e.com'>link</a></p>"
            "<ul><li>one</li><li>two</li></ul></body></html>"
        )
        c = doc_import.import_document("page.html", html.encode())["content"]
        self.assertIn("# Hi", c)
        self.assertIn("**bold**", c)
        self.assertIn("[link](http://e.com)", c)
        self.assertIn("- one", c)
        self.assertNotIn("color:red", c)  # <style> stripped

    def test_docx_to_md(self):
        from docx import Document

        doc = Document()
        doc.add_heading("Heading One", level=1)
        doc.add_paragraph("a normal paragraph")
        doc.add_paragraph("bullet item", style="List Bullet")
        buf = io.BytesIO()
        doc.save(buf)
        c = doc_import.import_document("d.docx", buf.getvalue())["content"]
        self.assertIn("# Heading One", c)
        self.assertIn("a normal paragraph", c)
        self.assertIn("- bullet item", c)

    def test_unsupported_type(self):
        with self.assertRaises(ValueError):
            doc_import.import_document("a.xyz", b"data")

    def test_pdf_graceful_when_pypdf_missing(self):
        try:
            import pypdf  # noqa: F401

            return  # pypdf installed → nothing to assert here
        except Exception:
            pass
        with self.assertRaises(ValueError):
            doc_import.import_document("a.pdf", b"%PDF-1.4 not a real pdf")

    def test_html_em_italic_and_ol(self):
        html = "<p><em>slanted</em> and <i>also</i></p><ol><li>first</li><li>second</li></ol>"
        c = doc_import.import_document("x.html", html.encode())["content"]
        self.assertIn("*slanted*", c)
        self.assertIn("*also*", c)
        self.assertIn("1. first", c)
        self.assertIn("2. second", c)

    def test_name_strips_path_and_ext(self):
        r = doc_import.import_document("path/to/doc.txt", b"x")
        self.assertEqual(r["name"], "doc")

    def test_htm_extension_works(self):
        html = "<h2>Sub</h2><p>body</p>"
        c = doc_import.import_document("page.htm", html.encode())["content"]
        self.assertIn("## Sub", c)
        self.assertIn("body", c)

    def test_docx_heading2_and_numbered_list(self):
        from docx import Document

        doc = Document()
        doc.add_heading("Chapter Two", level=2)
        doc.add_paragraph("step one", style="List Number")
        buf = io.BytesIO()
        doc.save(buf)
        c = doc_import.import_document("doc.docx", buf.getvalue())["content"]
        self.assertIn("## Chapter Two", c)
        self.assertIn("1. step one", c)


if __name__ == "__main__":
    unittest.main()
