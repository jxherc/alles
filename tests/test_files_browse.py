import os
import tempfile
import time
from pathlib import Path
from unittest import mock

from services import files_store as fstore
from tests._client import ApiTest


class FilesBrowseTests(ApiTest):
    def setUp(self):
        super().setUp()
        self.tmp = tempfile.TemporaryDirectory()
        self.root = Path(self.tmp.name)
        self.fp = mock.patch.object(fstore, "files_dir", lambda: self.root)
        self.fp.start()

    def tearDown(self):
        self.fp.stop()
        self.tmp.cleanup()
        super().tearDown()

    def _w(self, rel, body, mtime=None):
        p = self.root / rel
        p.parent.mkdir(parents=True, exist_ok=True)
        if isinstance(body, bytes):
            p.write_bytes(body)
        else:
            p.write_text(body)
        if mtime is not None:
            os.utime(p, (mtime, mtime))
        return p

    # ---- duplicates ----
    def test_dup_groups_identical(self):
        self._w("a.txt", "same content here")
        self._w("sub/b.txt", "same content here")
        self._w("c.txt", "unrelated")
        groups = self.client.get("/api/files/duplicates").json()["groups"]
        self.assertEqual(len(groups), 1)
        self.assertEqual(sorted(groups[0]["paths"]), ["a.txt", "sub/b.txt"])

    def test_dup_group_carries_size(self):
        body = "x" * 42
        self._w("a.txt", body)
        self._w("b.txt", body)
        g = self.client.get("/api/files/duplicates").json()["groups"][0]
        self.assertEqual(g["size"], 42)

    def test_dup_ignores_singletons(self):
        self._w("a.txt", "one")
        self._w("b.txt", "two")
        self._w("c.txt", "three")
        self.assertEqual(self.client.get("/api/files/duplicates").json()["groups"], [])

    def test_dup_ignores_empty(self):
        # two empty files have identical (empty) content but must NOT be flagged
        self._w("e1.txt", "")
        self._w("e2.txt", "")
        self.assertEqual(self.client.get("/api/files/duplicates").json()["groups"], [])

    def test_dup_same_size_diff_content_not_grouped_and_unique_skipped(self):
        import hashlib

        self._w("s1.txt", "AAAA")  # size 4, unique content
        self._w("s2.txt", "BBBB")  # size 4, unique content (shares size with s1)
        self._w("t1.txt", "hello")  # size 5
        self._w("t2.txt", "hello")  # size 5 (a real content dup)
        self._w("lonely.txt", "a one-of-a-kind length")  # unique size -> never hashed
        real = hashlib.sha256
        hashed = []

        def spy(b=b""):
            hashed.append(bytes(b))
            return real(b)

        with mock.patch("hashlib.sha256", side_effect=spy):
            groups = self.client.get("/api/files/duplicates").json()["groups"]
        # correctness: only the identical pair groups; same-size-different-content does not
        self.assertEqual([sorted(g["paths"]) for g in groups], [["t1.txt", "t2.txt"]])
        # perf: the unique-size file is never read/hashed; the shared-size ones are
        self.assertNotIn(b"a one-of-a-kind length", hashed)
        self.assertIn(b"hello", hashed)

    def test_dup_sorted_by_group_size(self):
        # a 3-file group should sort before a 2-file group
        for n in ("a", "b", "c"):
            self._w(f"{n}.txt", "trio")
        self._w("d.txt", "duo")
        self._w("e.txt", "duo")
        groups = self.client.get("/api/files/duplicates").json()["groups"]
        self.assertEqual([len(g["paths"]) for g in groups], [3, 2])

    # ---- preview ----
    def test_preview_docx_text(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello from a docx")
        doc.add_paragraph("second paragraph")
        doc.save(str(self.root / "report.docx"))
        d = self.client.get("/api/files/preview?path=report.docx").json()
        self.assertEqual(d["kind"], "docx")
        self.assertIn("Hello from a docx", d["text"])
        self.assertIn("second paragraph", d["text"])

    def test_preview_txt(self):
        self._w("notes.txt", "plain text body")
        d = self.client.get("/api/files/preview?path=notes.txt").json()
        self.assertEqual(d["kind"], "text")
        self.assertEqual(d["text"], "plain text body")

    def test_preview_unsupported(self):
        self._w("blob.bin", b"\x00\x01\x02")
        self.assertEqual(
            self.client.get("/api/files/preview?path=blob.bin").json()["kind"], "unsupported"
        )

    def test_preview_missing_404(self):
        r = self.client.get("/api/files/preview?path=nope.txt")
        self.assertEqual(r.status_code, 404)

    def test_preview_xlsx_graceful(self):
        # openpyxl isn't installed in this env — must degrade gracefully, not crash
        self._w("sheet.xlsx", b"not really xlsx")
        d = self.client.get("/api/files/preview?path=sheet.xlsx").json()
        self.assertEqual(d["kind"], "xlsx")
        self.assertIn("error", d)

    def test_preview_xlsx_closes_workbook(self):
        # read-only openpyxl keeps the file handle open until .close(); preview must close it.
        # openpyxl isn't installed, so inject a fake module to exercise the real code path.
        import sys
        import types

        self._w("sheet.xlsx", b"ignored; load_workbook is mocked")
        wb = mock.MagicMock()
        wb.active.iter_rows.return_value = [("a", "b"), (1, None)]
        fake = types.ModuleType("openpyxl")
        fake.load_workbook = mock.MagicMock(return_value=wb)
        with mock.patch.dict(sys.modules, {"openpyxl": fake}):
            d = self.client.get("/api/files/preview?path=sheet.xlsx").json()
        self.assertEqual(d["kind"], "xlsx")
        self.assertEqual(d["rows"], [["a", "b"], ["1", ""]])  # None -> "", ints -> str
        wb.close.assert_called_once()

    # ---- activity ----
    def test_activity_recent_first(self):
        now = time.time()
        self._w("old.txt", "o", mtime=now - 5 * 86400)
        self._w("new.txt", "n", mtime=now - 1 * 86400)
        items = self.client.get("/api/files/activity").json()["items"]
        self.assertEqual([i["path"] for i in items], ["new.txt", "old.txt"])

    def test_activity_excludes_old(self):
        now = time.time()
        self._w("recent.txt", "r", mtime=now - 2 * 86400)
        self._w("ancient.txt", "a", mtime=now - 90 * 86400)
        items = self.client.get("/api/files/activity?days=30").json()["items"]
        self.assertEqual([i["path"] for i in items], ["recent.txt"])

    def test_activity_limit(self):
        now = time.time()
        for i in range(5):
            self._w(f"f{i}.txt", str(i), mtime=now - i * 3600)
        items = self.client.get("/api/files/activity?limit=2").json()["items"]
        self.assertEqual(len(items), 2)
